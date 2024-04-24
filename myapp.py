from docx import Document
from docx.shared import Inches
import pyttsx3 # this one is for text tp speech

def speak(text):
    pyttsx3.speak(text)

document = Document()

# profile picture
document.add_picture(
    'Profile.png',
    width = Inches(2.0)
)

# name phone and email
name = input('What is your name? ')
speak('Hello ' + name + 'How are you today? ')

speak ('What is your phone number? ')
phone = input('What is your phone number ')
email = input('What is your email? ')

document.add_paragraph('Name - ' + name + ' | ' + 'Phone - ' + phone + ' | ' + 'email - ' + email)

# about.me
document.add_heading('About me')
document.add_paragraph(input('Tell about yourself '))

# Skills and Experience
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter Company ')
start_date = input('Enter Start Date ')
end_date = input('Enter End Date ')
p.add_run(company + ' ').bold = True
p.add_run(start_date + '-' + end_date + '\n').italic = True
experience_details = input(
    'Describe your experience at ' + company + ' '
)
p.add_run(experience_details)

#more experience
while True:
    has_more_experiences = input(
        'Do you have more expiriences ? Yes or No '
        )
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter Company ')
        start_date = input('Enter Start Date ')
        end_date = input('Enter End Date ')
        p.add_run(company + ' ').bold = True
        p.add_run(start_date + '-' + end_date + '\n').italic = True
        experience_details = input(
                'Describe your experience at ' + company + ' '
        )
        p.add_run(experience_details)
    else :
        break

# Skills
document.add_heading('Skills')
skills = input('Enter your skill ')
p = document.add_paragraph(skills)
p.style = 'List Bullet'

while True:
    has_more_skills = input('Do you have more skills to add? Yes or No ')
    if has_more_skills.lower() == 'yes':
        skills = input('Enter your skill ')
        p = document.add_paragraph(skills)
        p.style = 'List Bullet'  
    else :
        break

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text ="CV generated using python with help of amigoscode"
document.save('cv.docx')